#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2020/3/12 21:31
# @Author  : Roohom
# @Site    : 
# @File    : word1.py
# @Software: PyCharm
"""
读取word文件
"""
from docx import Document

doc = Document('test.docx') # 该程序列表下有一个test.docx的文件
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)

print(''.join(content))


# 示例程序，详情可见：https://www.jianshu.com/p/94ac13f6633e

"""
# coding:utf-8
# 写word文档文件
import sys

from docx import Document
from docx.shared import Inches

def main():
    reload(sys)
    sys.setdefaultencoding('utf-8')
    
    # 创建文档对象
    document = Document()
    
    # 设置文档标题，中文要用unicode字符串
    document.add_heading(u'我的一个新文档',0)
    
    # 往文档中添加段落
    p = document.add_paragraph('This is a paragraph having some ')
    p.add_run('bold ').bold = True
    p.add_run('and some ')
    p.add_run('italic.').italic = True
    
    # 添加一级标题
    document.add_heading(u'一级标题, level = 1',level = 1)
    document.add_paragraph('Intense quote',style = 'IntenseQuote')
    
    # 添加无序列表
    document.add_paragraph('first item in unordered list',style = 'ListBullet')
    
    # 添加有序列表
    document.add_paragraph('first item in ordered list',style = 'ListNumber')
    document.add_paragraph('second item in ordered list',style = 'ListNumber')
    document.add_paragraph('third item in ordered list',style = 'ListNumber')
    
    # 添加图片，并指定宽度
    document.add_picture('e:/docs/pic.png',width = Inches(1.25))
    
    # 添加表格: 1行3列
    table = document.add_table(rows = 1,cols = 3)
    # 获取第一行的单元格列表对象
    hdr_cells = table.rows[0].cells
    # 为每一个单元格赋值
    # 注：值都要为字符串类型
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'Tel'
    # 为表格添加一行
    new_cells = table.add_row().cells
    new_cells[0].text = 'Tom'
    new_cells[1].text = '19'
    new_cells[2].text = '12345678'
    
    # 添加分页符
    document.add_page_break()
    
    # 往新的一页中添加段落
    p = document.add_paragraph('This is a paragraph in new page.')
    
    # 保存文档
    document.save('e:/docs/demo1.docx')
    
if __name__ == '__main__':
    main()


# coding:utf-8
# 读取已有的word文档
import sys

from docx import Document

def main():
    reload(sys)
    sys.setdefaultencoding('utf-8')
    
    # 创建文档对象
    document = Document('e:/docs/demo2.docx')
    
    # 读取文档中所有的段落列表
    ps = document.paragraphs
    # 每个段落有两个属性：style和text
    ps_detail = [(x.text,x.style.name) for x in ps]
    with open('out.tmp','w+') as fout:
        fout.write('')
    # 读取段落并写入一个文件
    with open('out.tmp','a+') as fout:
        for p in ps_detail:
            fout.write(p[0] + '\t' + p[1] + '\n\n')
    
    # 读取文档中的所有段落的列表
    tables = document.tables
    # 遍历table，并将所有单元格内容写入文件中
    with open('out.tmp','a+') as fout:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    fout.write(cell.text + '\t')
                fout.write('\n')
    
if __name__ == '__main__':
    main()











"""